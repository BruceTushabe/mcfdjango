from django.http import HttpResponse, FileResponse
from django.shortcuts import render
from django.template.loader import render_to_string
from django.http import FileResponse
from django.utils.http import urlquote
from openpyxl import load_workbook
from docx import Document
import os

# Create your views here.

def index(request):
    return render(request, 'core/index.html')

def contact(request):
    return render(request, 'core/contact.html')

def upload(request):
    return render(request, 'core/upload.html')

def form(request):
    return render(request, 'core/form.html')

# def download(request):
  #  return render(request, 'core/download.html')


# Function to find data for a specific account number in the Excel sheet
def find_account_data(account_number):
    account_number = int(account_number) # Convert to string for comparison
    workbook = load_workbook('data1.xlsx', data_only=True)
    sheet = workbook.active
    for row in sheet.iter_rows(min_row=2, values_only=True):
    #    print(f"Checking row: {row}")  # Print the entire row for debugging
        if row[2] == account_number:  # Assuming the account number is in the 6th column (index 4)
            return row  # Return the row data if account number is found
    return None  # Return None if account number is not found

# Function to populate the Word document with data for a specific account number
def populate_document_for_account(document, account_data):
    if account_data:
        Address = account_data[21]
        ACCT_NAME = account_data[3]
        DIS_AMT = account_data[7]
        DIS_SHDL_DATE = account_data[8]
        AGE = account_data[16]
        Gender = account_data[20]
        DATE_ARREARS_START = account_data[10]
        Application_date = account_data[19]
        DOB = account_data[18]
        AMOUNT_CLAIMED = account_data[11]
        ARREARSDAYS = account_data[4]
        TERM = account_data[6]

        # Convert values to strings
        Address = str(Address)
        ACCT_NAME = str(ACCT_NAME)
        DIS_AMT = str(DIS_AMT)
        DIS_SHDL_DATE = str(DIS_SHDL_DATE)
        AGE = str(AGE)
        Gender = str(Gender)
        DATE_ARREARS_START = str(DATE_ARREARS_START)
        Application_date = str(Application_date)
        DOB = str(DOB)
        AMOUNT_CLAIMED = str(AMOUNT_CLAIMED)
        ARREARSDAYS = str(ARREARSDAYS)
        TERM = str(TERM)

        # Assuming the template has placeholders like "{{Name}}", "{{disAmount}}", "{{disDate}}"
        for paragraph in document.paragraphs:
            if "{{Name}}" in paragraph.text:
                paragraph.text = paragraph.text.replace("{{Name}}", ACCT_NAME)
            if "{{disAmount}}" in paragraph.text:
                paragraph.text = paragraph.text.replace("{{disAmount}}", str(DIS_AMT))
            if "{{disDate}}" in paragraph.text:
                paragraph.text = paragraph.text.replace("{{disDate}}", str(DIS_SHDL_DATE))
            if "{{address}}" in paragraph.text:
                paragraph.text = paragraph.text.replace("{{address}}", str(Address))
            if "{{age}}" in paragraph.text:
                paragraph.text = paragraph.text.replace("{{age}}", str(AGE))
            if "{{gender}}" in paragraph.text:
                paragraph.text = paragraph.text.replace("{{gender}}", str(Gender))
            if "{{dob}}" in paragraph.text:
                paragraph.text = paragraph.text.replace("{{dob}}", str(DOB))
            if "{{appdate}}" in paragraph.text:
                paragraph.text = paragraph.text.replace("{{appdate}}", str(Application_date))
            if "{{amtclaimed}}" in paragraph.text:
                paragraph.text = paragraph.text.replace("{{amtclaimed}}", str(AMOUNT_CLAIMED))
            if "{{datearrears}}" in paragraph.text:
                paragraph.text = paragraph.text.replace("{{datearreas}}", str(DATE_ARREARS_START))
            if "{{arrearsdays}}" in paragraph.text:
                paragraph.text = paragraph.text.replace("{{arrearsdays}}", str(ARREARSDAYS))
            if "{{loanterm}}" in paragraph.text:
                paragraph.text = paragraph.text.replace("{{loanterm}}", str(TERM))

        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    if "{{Name}}" in cell.text:
                        cell.text = cell.text.replace("{{Name}}", ACCT_NAME)
                    if "{{disAmount}}" in cell.text:
                        cell.text = cell.text.replace("{{disAmount}}", str(DIS_AMT))
                    if "{{disDate}}" in cell.text:
                        cell.text = cell.text.replace("{{disDate}}", str(DIS_SHDL_DATE))
                    if "{{address}}" in cell.text:
                        cell.text = cell.text.replace("{{address}}", str(Address))
                    if "{{age}}" in cell.text:
                        cell.text = cell.text.replace("{{age}}", str(AGE))
                    if "{{gender}}" in cell.text:
                        cell.text = cell.text.replace("{{gender}}", str(Gender))
                    if "{{dob}}" in cell.text:
                        cell.text = cell.text.replace("{{dob}}", str(DOB))
                    if "{{appdate}}" in cell.text:
                        cell.text = cell.text.replace("{{appdate}}", str(Application_date))
                    if "{{amtclaimed}}" in cell.text:
                        cell.text = cell.text.replace("{{amtclaimed}}", str(AMOUNT_CLAIMED))
                    if "{{datearrears}}" in cell.text:
                        cell.text = cell.text.replace("{{datearreas}}", str(DATE_ARREARS_START))
                    if "{{arrearsdays}}" in cell.text:
                        cell.text = cell.text.replace("{{arrearsdays}}", str(ARREARSDAYS))
                    if "{{loanterm}}" in cell.text:
                        cell.text = cell.text.replace("{{loanterm}}", str(TERM))
        return True  # Data populated successfully
    return False  # Account number not found

def generate(request): # for generating word documents
    if request.method == 'POST':
        account_number = request.POST.get('account_number')
        account_data = find_account_data(account_number)
        if account_data:
            document = Document('template.docx')
            populate_document_for_account(document, account_data)
            output_filename = f"output_{account_number}.docx"
            document.save(output_filename)
            return download(request, account_number)
        else:
            return HttpResponse("Account number not found")
    else:
        return render(request, 'core/generate.html')
    

# Load the Word document template
document = Document('template.docx')

def download(request, account_number):
    # Assuming the generated Word document is stored in the same directory as the Django app
    file_path = os.path.join(os.path.dirname(__file__), f"output_{account_number}.docx")
    
    # Check if the file exists
    if os.path.exists(file_path):
        # Open the file in binary mode and serve it as a downloadable attachment
        with open(file_path, 'rb') as f:
            response = FileResponse(f)
            response['Content-Disposition'] = f'attachment; filename="output_{account_number}.docx"'
            return response
    else:
        # If the file does not exist, return an error message
        return HttpResponse("File not found", status=404)